#!/usr/bin/env python3
"""
文件格式转换脚本

将 PDF、DOCX、DOC、TXT、HTML、视频链接、音频链接转换为 Markdown 格式，
放入 raw/ 目录供 LLM Wiki 的 Ingest 流程使用。

用法：
  python3 scripts/convert-to-raw.py <输入文件或目录或链接> [--topic <topic>] [--output-dir <目录>]

示例：
  # 转换文件
  python3 scripts/convert-to-raw.py paper.pdf --topic research

  # 转换视频链接
  python3 scripts/convert-to-raw.py "https://www.youtube.com/watch?v=xxx" --topic ai

  # 转换音频链接
  python3 scripts/convert-to-raw.py "https://www.ximalaya.com/xxx" --topic podcast

  # 转换整个目录
  python3 scripts/convert-to-raw.py ./downloads/ --topic work
"""

import os
import re
import sys
import argparse
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime
from urllib.parse import urlparse

# ============================================================
# 配置
# ============================================================
DEFAULT_RAW_DIR = Path(__file__).parent.parent / "raw"
DEFAULT_TOPIC = "general"

SUPPORTED_EXTENSIONS = {
    '.pdf': 'PDF',
    '.docx': 'Word (DOCX)',
    '.doc': 'Word (DOC)',
    '.txt': '纯文本 (TXT)',
    '.html': 'HTML',
    '.htm': 'HTML',
    '.md': 'Markdown',
    '.markdown': 'Markdown',
    '.mp3': '音频 (MP3)',
    '.mp4': '视频 (MP4)',
    '.m4a': '音频 (M4A)',
    '.wav': '音频 (WAV)',
    '.flac': '音频 (FLAC)',
    '.mkv': '视频 (MKV)',
    '.webm': '视频 (WebM)',
}

# 视频平台域名映射
VIDEO_PLATFORMS = {
    'youtube.com', 'youtu.be', 'www.youtube.com', 'm.youtube.com',
    'bilibili.com', 'www.bilibili.com', 'm.bilibili.com',
    'vimeo.com', 'www.vimeo.com',
    'douyin.com', 'www.douyin.com',
    'tiktok.com', 'www.tiktok.com',
    'ixigua.com', 'www.ixigua.com',
    'zhihu.com', 'www.zhihu.com',
    'ximalaya.com', 'www.ximalaya.com',
}

# ============================================================
# 转换器
# ============================================================

def convert_txt(file_path: Path) -> str:
    """TXT 文件直接读取，清理格式"""
    content = file_path.read_text(encoding='utf-8', errors='replace')
    # 清理多余空行
    content = re.sub(r'\n{3,}', '\n\n', content)
    return content.strip()


def convert_pdf(file_path: Path) -> str:
    """PDF 转 Markdown"""
    try:
        import pdfplumber
    except ImportError:
        print("  ⚠️  需要安装 pdfplumber：pip install pdfplumber --break-system-packages")
        sys.exit(1)

    text_parts = []
    with pdfplumber.open(str(file_path)) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"<!-- 第 {i+1} 页 -->\n\n{text}")
    
    if not text_parts:
        return ""
    
    return "\n\n---\n\n".join(text_parts)


def convert_docx(file_path: Path) -> str:
    """DOCX 转 Markdown"""
    try:
        from docx import Document
    except ImportError:
        print("  ⚠️  需要安装 python-docx：pip install python-docx --break-system-packages")
        sys.exit(1)

    doc = Document(str(file_path))
    md_lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("")
            continue
        
        # 检测标题样式
        style_name = para.style.name if para.style else ""
        if "Heading 1" in style_name or "标题 1" in style_name:
            md_lines.append(f"# {text}")
        elif "Heading 2" in style_name or "标题 2" in style_name:
            md_lines.append(f"## {text}")
        elif "Heading 3" in style_name or "标题 3" in style_name:
            md_lines.append(f"### {text}")
        elif "Heading 4" in style_name or "标题 4" in style_name:
            md_lines.append(f"#### {text}")
        else:
            md_lines.append(text)
    
    # 处理表格
    for table in doc.tables:
        md_lines.append("")
        # 表头
        header = [cell.text.strip() for cell in table.rows[0].cells]
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        # 表体
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("")
    
    return "\n".join(md_lines)


def convert_doc(file_path: Path) -> str:
    """DOC 转 Markdown（通过 LibreOffice 或 antiword）"""
    import subprocess
    import tempfile
    
    # 尝试使用 LibreOffice 转换为 DOCX，再用 docx 转换器
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'docx',
                 '--outdir', tmpdir, str(file_path)],
                capture_output=True, timeout=60
            )
            docx_file = Path(tmpdir) / (file_path.stem + ".docx")
            if docx_file.exists():
                return convert_docx(docx_file)
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    
    # 尝试使用 antiword
    try:
        result = subprocess.run(
            ['antiword', str(file_path)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    
    print("  ⚠️  无法转换 .doc 文件。请安装以下工具之一：")
    print("     - LibreOffice: https://www.libreoffice.org/")
    print("     - antiword: sudo apt install antiword")
    print("     - 或者手动将 .doc 另存为 .docx 后重试")
    return ""


def convert_html(file_path: Path) -> str:
    """HTML 转 Markdown"""
    try:
        from html.parser import HTMLParser
    except ImportError:
        pass
    
    # 简单的 HTML → 文本转换
    content = file_path.read_text(encoding='utf-8', errors='replace')
    
    # 移除 script 和 style
    content = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL | re.IGNORECASE)
    
    # 转换常见标签
    content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'#### \1', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'[\2](\1)', content, flags=re.DOTALL | re.IGNORECASE)
    content = re.sub(r'<img[^>]*src="([^"]*)"[^>]*alt="([^"]*)"[^>]*/?>', r'![\2](\1)', content, flags=re.IGNORECASE)
    
    # 移除所有剩余标签
    content = re.sub(r'<[^>]+>', '', content)
    
    # 解码 HTML 实体
    content = content.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&quot;', '"').replace('&#39;', "'")
    
    # 清理
    content = re.sub(r'\n{3,}', '\n\n', content)
    return content.strip()


def convert_markdown(file_path: Path) -> str:
    """Markdown 文件直接读取"""
    return file_path.read_text(encoding='utf-8', errors='replace')


def is_url(text: str) -> bool:
    """判断输入是否为 URL"""
    return text.strip().startswith('http://') or text.strip().startswith('https://')


def is_video_url(url: str) -> bool:
    """判断是否为视频平台 URL"""
    try:
        domain = urlparse(url).hostname or ""
        return domain in VIDEO_PLATFORMS
    except Exception:
        return False


def convert_video_url(url: str, topic: str, output_dir: Path) -> Path:
    """将视频链接转换为 Markdown（通过 yt-dlp 提取字幕/转录）"""
    # 检查 yt-dlp
    try:
        subprocess.run(['yt-dlp', '--version'], capture_output=True, check=True)
    except (FileNotFoundError, subprocess.CalledProcessError):
        print("  ⚠️  需要安装 yt-dlp：")
        print("     pip install yt-dlp --break-system-packages")
        print("     或: brew install yt-dlp (macOS)")
        return None

    print(f"  🎬 处理视频链接: {url[:60]}...")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)

        # 策略 1：尝试下载字幕（最快）
        print("  📥 尝试提取字幕...")
        subtitle_file = tmpdir_path / "subtitle"
        try:
            subprocess.run(
                ['yt-dlp', '--write-sub', '--write-auto-sub',
                 '--sub-lang', 'zh,en',
                 '--skip-download',
                 '--convert-subs', 'srt',
                 '-o', str(subtitle_file),
                 url],
                capture_output=True, timeout=120
            )
        except subprocess.TimeoutExpired:
            print("  ⚠️  字幕提取超时")

        # 查找下载的字幕文件
        srt_files = list(tmpdir_path.glob("*.srt"))
        if srt_files:
            srt_content = srt_files[0].read_text(encoding='utf-8', errors='replace')
            content = _srt_to_markdown(srt_content)
            if content:
                print("  ✅ 字幕提取成功")
                return _save_converted_content(url, content, topic, output_dir, "video", url)

        # 策略 2：尝试下载音频并用 Whisper 转录
        print("  🎤 字幕不可用，尝试音频转录...")
        audio_file = tmpdir_path / "audio.mp3"
        try:
            subprocess.run(
                ['yt-dlp', '-x', '--audio-format', 'mp3',
                 '--audio-quality', '5',
                 '-o', str(audio_file),
                 url],
                capture_output=True, timeout=300
            )
        except subprocess.TimeoutExpired:
            print("  ⚠️  音频下载超时")
            return None

        if not audio_file.exists():
            print("  ⚠️  音频下载失败")
            return None

        # 使用 Whisper 转录
        content = _transcribe_audio(audio_file)
        if content:
            print("  ✅ 音频转录成功")
            return _save_converted_content(url, content, topic, output_dir, "video", url)

    print("  ⚠️  所有转换策略均失败")
    print("     建议：手动下载字幕文件（.srt）后放入 raw/ 目录")
    return None


def convert_audio_file(file_path: Path) -> str:
    """将本地音频文件转换为 Markdown（通过 Whisper 转录）"""
    content = _transcribe_audio(file_path)
    if not content:
        print("  ⚠️  音频转录失败")
    return content


def convert_video_file(file_path: Path) -> str:
    """将本地视频文件转换为 Markdown（提取音频后 Whisper 转录）"""
    # 先提取音频
    try:
        subprocess.run(['ffmpeg', '-version'], capture_output=True, check=True)
    except (FileNotFoundError, subprocess.CalledProcessError):
        print("  ⚠️  需要安装 ffmpeg：")
        print("     Ubuntu: sudo apt install ffmpeg")
        print("     macOS: brew install ffmpeg")
        return ""

    with tempfile.TemporaryDirectory() as tmpdir:
        audio_file = Path(tmpdir) / "audio.mp3"
        try:
            subprocess.run(
                ['ffmpeg', '-i', str(file_path), '-vn', '-acodec', 'libmp3lame',
                 '-q:a', '5', '-y', str(audio_file)],
                capture_output=True, timeout=300
            )
        except subprocess.TimeoutExpired:
            print("  ⚠️  音频提取超时")
            return ""

        if audio_file.exists():
            return _transcribe_audio(audio_file) or ""
        return ""


def _srt_to_markdown(srt_content: str) -> str:
    """将 SRT 字幕格式转换为 Markdown"""
    lines = []
    for block in re.split(r'\n\n+', srt_content.strip()):
        block_lines = block.strip().split('\n')
        if len(block_lines) < 2:
            continue
        # 跳过序号行，提取时间戳和文本
        for line in block_lines[1:]:
            line = line.strip()
            if not line:
                continue
            # 跳过纯时间戳行
            if re.match(r'^\d{2}:\d{2}:\d{2}', line) and '-->' in line:
                continue
            lines.append(line)

    return '\n\n'.join(lines).strip()


def _transcribe_audio(audio_path: Path) -> str:
    """使用 Whisper 转录音频文件"""
    try:
        import whisper
    except ImportError:
        print("  ⚠️  需要安装 whisper：")
        print("     pip install openai-whisper --break-system-packages")
        print("     或使用更轻量的: pip install faster-whisper --break-system-packages")
        return ""

    try:
        model = whisper.load_model("base")  # base 模型约 74MB，速度较快
        result = model.transcribe(str(audio_path), language="zh", verbose=False)
        return result["text"].strip()
    except Exception as e:
        print(f"  ⚠️  Whisper 转录失败: {e}")
        return ""


def _save_converted_content(url: str, content: str, topic: str,
                            output_dir: Path, original_format: str,
                            source_url: str) -> Path:
    """保存转换后的内容为 Markdown 文件"""
    if not content:
        return None

    # 从 URL 提取标题
    slug = re.sub(r'[^\w\u4e00-\u9fff-]+', '-', url.split('/')[-1])[:60]
    today = datetime.now().strftime("%Y-%m-%d")
    output_name = f"{today}-{slug}.md"
    output_path = output_dir / topic / output_name

    counter = 2
    while output_path.exists():
        output_name = f"{today}-{slug}-{counter}.md"
        output_path = output_dir / topic / output_name
        counter += 1

    frontmatter = f"""---
source: {source_url}
collected: {today}
published: Unknown
original_format: {original_format}
---

"""
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path.write_text(frontmatter + content, encoding='utf-8')
    print(f"  ✅ 已保存: raw/{topic}/{output_name}")
    return output_path


# ============================================================
# 转换调度
# ============================================================

CONVERTERS = {
    '.txt': convert_txt,
    '.pdf': convert_pdf,
    '.docx': convert_docx,
    '.doc': convert_doc,
    '.html': convert_html,
    '.htm': convert_html,
    '.md': convert_markdown,
    '.markdown': convert_markdown,
    '.mp3': convert_audio_file,
    '.m4a': convert_audio_file,
    '.wav': convert_audio_file,
    '.flac': convert_audio_file,
    '.mp4': convert_video_file,
    '.mkv': convert_video_file,
    '.webm': convert_video_file,
}


def convert_file(file_path: Path, topic: str, output_dir: Path) -> Path:
    """转换单个文件为 Markdown 并保存到 raw/<topic>/"""
    ext = file_path.suffix.lower()
    
    if ext not in CONVERTERS:
        print(f"  ⚠️  跳过不支持的格式: {file_path.name} ({ext})")
        return None
    
    print(f"  📄 转换 {file_path.name} ({SUPPORTED_EXTENSIONS[ext]})...")
    
    # 执行转换
    content = CONVERTERS[ext](file_path)
    
    if not content:
        print(f"  ⚠️  转换结果为空，跳过: {file_path.name}")
        return None
    
    # 生成输出文件
    slug = re.sub(r'[^\w\u4e00-\u9fff-]+', '-', file_path.stem)[:60]
    today = datetime.now().strftime("%Y-%m-%d")
    output_name = f"{today}-{slug}.md"
    output_path = output_dir / topic / output_name
    
    # 避免文件名冲突
    counter = 2
    while output_path.exists():
        output_name = f"{today}-{slug}-{counter}.md"
        output_path = output_dir / topic / output_name
        counter += 1
    
    # 写入文件（带元数据）
    source_url = ""
    if ext == '.html' or ext == '.htm':
        source_url = file_path.as_uri()
    
    frontmatter = f"""---
source: {source_url or file_path.name}
collected: {today}
published: Unknown
original_format: {ext}
---

"""
    
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path.write_text(frontmatter + content, encoding='utf-8')
    
    print(f"  ✅ 已保存: raw/{topic}/{output_name}")
    return output_path


def convert_directory(dir_path: Path, topic: str, output_dir: Path):
    """转换目录中的所有支持的文件"""
    converted = []
    skipped = []
    
    for file_path in sorted(dir_path.rglob("*")):
        if not file_path.is_file():
            continue
        ext = file_path.suffix.lower()
        if ext in CONVERTERS:
            result = convert_file(file_path, topic, output_dir)
            if result:
                converted.append(result)
        else:
            skipped.append(file_path.name)
    
    return converted, skipped


# ============================================================
# 主函数
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="将文件/视频链接/音频链接转换为 Markdown 格式，放入 raw/ 目录",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
支持的格式：
  文件格式：
    .pdf    PDF 文档
    .docx   Word 文档
    .doc    Word 文档（需要 LibreOffice 或 antiword）
    .txt    纯文本
    .html   网页
    .md     Markdown（直接复制）
    .mp3/.m4a/.wav/.flac  音频文件（需要 Whisper）
    .mp4/.mkv/.webm       视频文件（需要 ffmpeg + Whisper）

  链接格式：
    YouTube / Bilibili / 抖音 / TikTok / 西瓜视频  视频链接
    喜马拉雅 / 其他音频平台                     音频链接

依赖安装：
  基础：pip install pdfplumber python-docx yt-dlp --break-system-packages
  音视频转录：pip install openai-whisper --break-system-packages
  视频提取：sudo apt install ffmpeg

转换策略（视频/音频链接）：
  1. 优先提取平台字幕（最快，无需额外依赖）
  2. 字幕不可用时，下载音频用 Whisper 转录（需要 Whisper）

示例：
  python3 scripts/convert-to-raw.py paper.pdf --topic research
  python3 scripts/convert-to-raw.py "https://www.youtube.com/watch?v=xxx" --topic ai
  python3 scripts/convert-to-raw.py "https://www.bilibili.com/video/BV1xx" --topic tutorial
  python3 scripts/convert-to-raw.py podcast.mp3 --topic podcast
  python3 scripts/convert-to-raw.py ./downloads/ --topic work
"""
    )
    parser.add_argument("input", help="输入文件、目录或链接路径")
    parser.add_argument("--topic", "-t", default=DEFAULT_TOPIC, help=f"主题目录名（默认: {DEFAULT_TOPIC}）")
    parser.add_argument("--output-dir", "-o", default=str(DEFAULT_RAW_DIR), help="输出目录（默认: raw/）")
    
    args = parser.parse_args()
    
    input_str = args.input.strip()
    output_dir = Path(args.output_dir)
    
    print("=" * 50)
    print("🔄 文件格式转换工具")
    print("=" * 50)
    print(f"  输入: {input_str[:80]}")
    print(f"  主题: {args.topic}")
    print(f"  输出: {output_dir}/{args.topic}/")
    print()
    
    # 处理 URL 输入
    if is_url(input_str):
        if is_video_url(input_str):
            result = convert_video_url(input_str, args.topic, output_dir)
            if result:
                print(f"\n✅ 视频链接转换完成")
            else:
                print(f"\n❌ 视频链接转换失败")
                sys.exit(1)
        else:
            print(f"⚠️  该链接平台暂不支持自动转换")
            print(f"   建议：使用 Obsidian Web Clipper 保存网页内容")
            sys.exit(1)
        return
    
    # 处理文件/目录输入
    input_path = Path(input_str)
    
    if not input_path.exists():
        print(f"❌ 路径不存在: {input_path}")
        sys.exit(1)
    
    print("=" * 50)
    print("🔄 文件格式转换工具")
    print("=" * 50)
    print(f"  输入: {input_path}")
    print(f"  主题: {args.topic}")
    print(f"  输出: {output_dir}/{args.topic}/")
    print()
    
    if input_path.is_file():
        convert_file(input_path, args.topic, output_dir)
    elif input_path.is_dir():
        converted, skipped = convert_directory(input_path, args.topic, output_dir)
        print()
        print(f"✅ 转换完成: {len(converted)} 个文件")
        if skipped:
            print(f"⚠️  跳过: {len(skipped)} 个不支持的文件")
            for name in skipped[:10]:
                print(f"   - {name}")
            if len(skipped) > 10:
                print(f"   ... 还有 {len(skipped) - 10} 个")
    else:
        print(f"❌ 不支持的输入类型: {input_path}")
        sys.exit(1)


if __name__ == "__main__":
    main()
